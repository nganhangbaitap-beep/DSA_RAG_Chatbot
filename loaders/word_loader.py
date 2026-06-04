# loaders/word_loader.py
# Đọc file Word (.docx, .doc)

import os
from langchain_core.documents import Document


def load_word(file_path: str) -> list[Document]:
    """
    Đọc file Word (.docx).
    Trích xuất: text, tiêu đề, bảng, danh sách.
    
    Args:
        file_path: Đường dẫn file .docx
        
    Returns:
        Danh sách Document (chia theo section/heading)
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        print("❌ Thiếu python-docx. Chạy: pip install python-docx")
        return []
    
    docs = []
    filename = os.path.basename(file_path)
    
    print(f"📝 Đang đọc Word: {filename}")
    
    try:
        docx = DocxDocument(file_path)
        
        # ===== Thu thập toàn bộ nội dung =====
        sections = []
        current_section = {"heading": "Tổng quan", "content": [], "level": 0}
        
        for para in docx.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Kiểm tra heading
            if para.style.name.startswith("Heading"):
                # Lưu section hiện tại nếu có nội dung
                if current_section["content"]:
                    sections.append(current_section)
                
                # Lấy level của heading
                try:
                    level = int(para.style.name.split()[-1])
                except ValueError:
                    level = 1
                
                current_section = {
                    "heading": text,
                    "content": [],
                    "level": level
                }
            else:
                current_section["content"].append(text)
        
        # Lưu section cuối
        if current_section["content"]:
            sections.append(current_section)
        
        # ===== Đọc bảng =====
        table_texts = []
        for table_idx, table in enumerate(docx.tables):
            table_content = [f"[BẢNG {table_idx + 1}]"]
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                row_data = [cell for cell in row_data if cell]
                if row_data:
                    table_content.append(" | ".join(row_data))
            
            if len(table_content) > 1:
                table_texts.append("\n".join(table_content))
        
        # ===== Tạo Documents =====
        for i, section in enumerate(sections):
            content = section["heading"] + "\n\n"
            content += "\n".join(section["content"])
            
            if content.strip():
                doc = Document(
                    page_content=content.strip(),
                    metadata={
                        "source": filename,
                        "source_type": "word",
                        "section": section["heading"],
                        "section_index": i,
                        "heading_level": section["level"],
                        "file_path": file_path,
                    }
                )
                docs.append(doc)
        
        # Thêm bảng vào cuối
        for i, table_text in enumerate(table_texts):
            doc = Document(
                page_content=table_text,
                metadata={
                    "source": filename,
                    "source_type": "word_table",
                    "table_index": i,
                    "file_path": file_path,
                }
            )
            docs.append(doc)
        
        print(f"   ✅ Đọc thành công: {len(sections)} section, {len(table_texts)} bảng")
        
    except Exception as e:
        print(f"   ❌ Lỗi đọc Word: {e}")
    
    return docs
